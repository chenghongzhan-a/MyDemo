import os, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ---- Page setup ----
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = '\u5fae\u8f6f\u96c5\u9ed1'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '\u5fae\u8f6f\u96c5\u9ed1')

# ---- Helper functions ----
def set_font(run, name='\u5fae\u8f6f\u96c5\u9ed1', size=Pt(10.5), bold=False, color=None):
    run.font.name = name
    run.font.size = size
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = color

def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=Pt(14), bold=True, color=RGBColor(0x1A, 0x3C, 0x6E))
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr ' + nsdecls("w") + '>'
        '<w:bottom w:val="single" w:sz="4" w:space="1" w:color="1A3C6E"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run('\u2022 ' + text)
    set_font(run)
    return p

# =====================
# HEADER
# =====================
name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_p.paragraph_format.space_after = Pt(4)
run = name_p.add_run('\u5f20 \u5b87 \u8f69')
set_font(run, size=Pt(22), bold=True, color=RGBColor(0x1A, 0x3C, 0x6E))

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_p.paragraph_format.space_after = Pt(2)
run = info_p.add_run('Email: zhangyuxuan@example.com  |  Tel: 138-xxxx-xxxx  |  DOB: 2001.06')
set_font(run, size=Pt(10))

intent_p = doc.add_paragraph()
intent_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
intent_p.paragraph_format.space_after = Pt(2)
run = intent_p.add_run('\u6c42\u804c\u610f\u5411\uff1aUnity\u524d\u7aef\u5f00\u53d1\u5de5\u7a0b\u5e08  |  \u610f\u5411\u57ce\u5e02\uff1a\u4e0d\u9650  |  \u968f\u65f6\u5230\u5c97')
set_font(run, size=Pt(10), color=RGBColor(0x55, 0x55, 0x55))

# =====================
# Education
# =====================
add_section_title(doc, '\u6559\u80b2\u80cc\u666f')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.line_spacing = 1.5
run = p.add_run('XX\u5927\u5b66')
set_font(run, bold=True)
run = p.add_run('\u3000\u3000\u8ba1\u7b97\u673a\u79d1\u5b66\u4e0e\u6280\u672f / \u8f6f\u4ef6\u5de5\u7a0b\u3000\u3000\u672c\u79d1\u3000\u30002021.09 - 2025.06')
set_font(run)

add_bullet(doc, '\u4e3b\u4fee\u8bfe\u7a0b\uff1a\u6570\u636e\u7ed3\u6784\u4e0e\u7b97\u6cd5\u3001\u8ba1\u7b97\u673a\u7f51\u7edc\u3001\u64cd\u4f5c\u7cfb\u7edf\u3001\u8ba1\u7b97\u673a\u56fe\u5f62\u5b66\u3001\u9762\u5411\u5bf9\u8c61\u7a0b\u5e8f\u8bbe\u8ba1\u7b49')
add_bullet(doc, '\u82f1\u8bed\u516d\u7ea7\uff08CET-6\uff09\uff0c\u53ef\u6d41\u7545\u9605\u8bfb\u82f1\u6587\u6280\u672f\u6587\u6863')

# =====================
# Skills
# =====================
add_section_title(doc, '\u4e13\u4e1a\u6280\u80fd')
skills = [
    '\u719f\u7ec3\u638c\u63e1 C# \u7f16\u7a0b\u8bed\u8a00\uff0c\u719f\u6089\u9762\u5411\u5bf9\u8c61\u7f16\u7a0b\u3001\u8bbe\u8ba1\u6a21\u5f0f\uff0c\u5177\u5907\u826f\u597d\u7684\u4ee3\u7801\u89c4\u8303\u610f\u8bc6',
    '\u719f\u7ec3\u638c\u63e1 Unity \u5f15\u64ce\uff0c\u719f\u6089 UGUI \u754c\u9762\u7cfb\u7edf\uff0c\u80fd\u591f\u72ec\u7acb\u5b8c\u6210 UI \u642d\u5efa\u3001\u9002\u914d\u4e0e\u4ea4\u4e92\u903b\u8f91',
    '\u719f\u6089 Lua \u811a\u672c\u8bed\u8a00\uff0c\u4e86\u89e3 XLua / ToLua \u7b49\u70ed\u66f4\u65b0\u65b9\u6848\u5728 Unity \u9879\u76ee\u4e2d\u7684\u5e94\u7528',
    '\u719f\u6089 Unity \u7f51\u7edc\u901a\u4fe1\u5f00\u53d1\uff0c\u638c\u63e1 TCP/UDP \u534f\u8bae\u57fa\u7840\uff0c\u6709\u5b9e\u9645\u8054\u673a\u529f\u80fd\u5f00\u53d1\u7ecf\u9a8c',
    '\u4e86\u89e3 2D \u6e38\u620f\u5f00\u53d1\u6d41\u7a0b\uff0c\u5177\u5907 Tilemap\u3001Sprite\u3001\u7269\u7406\u78b0\u649e\u7b49 2D \u6a21\u5757\u7684\u4f7f\u7528\u7ecf\u9a8c',
    '\u719f\u6089 Git \u7248\u672c\u63a7\u5236\uff0c\u5177\u5907\u56e2\u961f\u534f\u4f5c\u5f00\u53d1\u7684\u57fa\u672c\u80fd\u529b',
]
for s in skills:
    add_bullet(doc, s)

# =====================
# Project
# =====================
add_section_title(doc, '\u9879\u76ee\u7ecf\u5386')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(2)
run = p.add_run('2D \u4fef\u89c6\u89d2\u6c99\u76d2\u6e38\u620f Demo')
set_font(run, size=Pt(12), bold=True)
run = p.add_run('\u3000\u3000\u72ec\u7acb\u5f00\u53d1\u3000\u30002024.06 - 2024.10')
set_font(run, size=Pt(10), color=RGBColor(0x88, 0x88, 0x88))

project_points = [
    '\u57fa\u4e8e Unity \u5f15\u64ce\u5f00\u53d1\u7684\u4e00\u6b3e 2D \u4fef\u89c6\u89d2\u6c99\u76d2\u6e38\u620f\uff0c\u73a9\u5bb6\u53ef\u5728\u5f00\u653e\u4e16\u754c\u4e2d\u8fdb\u884c\u8d44\u6e90\u91c7\u96c6\u3001\u5efa\u9020\u4e0e\u63a2\u7d22',
    '\u4f7f\u7528 UGUI \u642d\u5efa\u5b8c\u6574\u7684\u6e38\u620f\u754c\u9762\u7cfb\u7edf\uff0c\u5305\u62ec\u80cc\u5305\u3001\u5feb\u6377\u680f\u3001\u5efa\u9020\u83dc\u5355\u3001\u8bbe\u7f6e\u9762\u677f\u7b49\u6a21\u5757\uff0c\u9002\u914d\u591a\u79cd\u5206\u8fa8\u7387',
    '\u91c7\u7528 Tilemap + Rule Tile \u5b9e\u73b0\u52a8\u6001\u5730\u5f62\u751f\u6210\u4e0e\u7ba1\u7406\uff0c\u652f\u6301\u73a9\u5bb6\u81ea\u7531\u6539\u9020\u5730\u5f62',
    '\u57fa\u4e8e Unity Netcode / UNET \u5b9e\u73b0\u5c40\u57df\u7f51\u591a\u4eba\u8054\u673a\u529f\u80fd\uff0c\u5305\u62ec\u72b6\u6001\u540c\u6b65\u3001RPC \u901a\u4fe1\u4e0e\u623f\u95f4\u7ba1\u7406',
    '\u5f15\u5165 Lua\uff08XLua\uff09\u5b9e\u73b0\u90e8\u5206\u6e38\u620f\u903b\u8f91\u7684\u70ed\u66f4\u65b0\uff0c\u964d\u4f4e\u7248\u672c\u8fed\u4ee3\u7684\u53d1\u5e03\u6210\u672c',
    '\u4f7f\u7528\u5bf9\u8c61\u6c60\u3001\u5f02\u6b65\u52a0\u8f7d\u7b49\u6280\u672f\u4f18\u5316\u5185\u5b58\u4e0e\u52a0\u8f7d\u6027\u80fd\uff0c\u786e\u4fdd\u4e2d\u4f4e\u7aef\u8bbe\u5907\u6d41\u7545\u8fd0\u884c',
    '\u9879\u76ee\u4ee3\u7801\u6258\u7ba1\u4e8e GitHub\uff0c\u5305\u542b\u5b8c\u6574\u7684\u6280\u672f\u6587\u6863\u4e0e\u4f7f\u7528\u8bf4\u660e',
]
for pt in project_points:
    add_bullet(doc, pt)

# =====================
# Other experience
# =====================
add_section_title(doc, '\u5176\u4ed6\u7ecf\u5386')
add_bullet(doc, '\u5728\u6821\u671f\u95f4\u53c2\u4e0e ACM \u6821\u8d5b\u5e76\u83b7\u5f97\u4e09\u7b49\u5956\uff0c\u5177\u5907\u8f83\u5f3a\u7684\u7b97\u6cd5\u4e0e\u903b\u8f91\u601d\u7ef4\u80fd\u529b')
add_bullet(doc, 'GitHub \u4e2a\u4eba\u6280\u672f\u535a\u5ba2 / \u6280\u672f\u7b14\u8bb0\uff0c\u6301\u7eed\u8f93\u51fa Unity \u4e0e C# \u76f8\u5173\u5b66\u4e60\u7b14\u8bb0\u4e0e\u6280\u672f\u5206\u4eab')

# =====================
# Self
# =====================
add_section_title(doc, '\u81ea\u6211\u8bc4\u4ef7')
self_text = (
    '\u70ed\u7231\u6e38\u620f\u5f00\u53d1\uff0c\u5bf9 Unity \u524d\u7aef\u5de5\u7a0b\u6709\u6d53\u539a\u5174\u8da3\u548c\u6301\u7eed\u5b66\u4e60\u7684\u70ed\u60c5\u3002'
    '\u81ea\u5b66\u80fd\u529b\u5f3a\uff0c\u5728\u5927\u5b66\u671f\u95f4\u901a\u8fc7\u72ec\u7acb\u5f00\u53d1\u5b8c\u6210\u5b8c\u6574\u6e38\u620f Demo\uff0c'
    '\u4ece\u9700\u6c42\u5206\u6790\u3001\u6280\u672f\u9009\u578b\u5230\u7f16\u7801\u5b9e\u73b0\u5168\u7a0b\u72ec\u7acb\u63a8\u52a8\u3002'
    '\u4ee3\u7801\u98ce\u683c\u4e25\u8c28\uff0c\u6ce8\u91cd\u53ef\u8bfb\u6027\u548c\u53ef\u7ef4\u62a4\u6027\u3002'
    '\u5584\u4e8e\u901a\u8fc7\u5b98\u65b9\u6587\u6863\u3001\u793e\u533a\u8d44\u6e90\u89e3\u51b3\u6280\u672f\u95ee\u9898\u3002'
    '\u6027\u683c\u79ef\u6781\u4e50\u89c2\uff0c\u6c9f\u901a\u8868\u8fbe\u6e05\u6670\uff0c\u671f\u5f85\u52a0\u5165\u4f18\u79c0\u7684\u7814\u53d1\u56e2\u961f\uff0c'
    '\u5728\u5b9e\u9645\u9879\u76ee\u4e2d\u5feb\u901f\u6210\u957f\u5e76\u4e3a\u4ea7\u54c1\u521b\u9020\u4ef7\u503c\u3002'
)
p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.6
p.paragraph_format.left_indent = Cm(0.5)
run = p.add_run(self_text)
set_font(run)

# ---- Save ----
out_dir = os.path.dirname(os.path.abspath(__file__))
save_path = os.path.join(out_dir, 'resume.docx')
doc.save(save_path)
print('OK: ' + save_path)
